import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_repeat_table_header(row):
    """ set repeat table header on every new page """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row

def add_header(doc, left_logo_path, right_logo_path):
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        
        # We use a table for the header to position logos left and right
        table = header.add_table(1, 2, doc.width if hasattr(doc, 'width') else Inches(6.5))
        table.width = Inches(6.5)
        
        # Left Logo
        cell_left = table.rows[0].cells[0]
        run_left = cell_left.paragraphs[0].add_run()
        if os.path.exists(left_logo_path):
            run_left.add_picture(left_logo_path, width=Inches(1.5))
        
        # Right Logo
        cell_right = table.rows[0].cells[1]
        p_right = cell_right.paragraphs[0]
        p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = p_right.add_run()
        if os.path.exists(right_logo_path):
            run_right.add_picture(right_logo_path, width=Inches(1.5))

def add_page_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 78, 121) # Dark Blue
    
    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(16)
        run2.font.bold = True
        run2.font.color.rgb = RGBColor(46, 116, 181) # Medium Blue

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

def add_image_with_caption(doc, img_path, caption):
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.italic = True
        p.style.font.size = Pt(10)
    else:
        doc.add_paragraph(f"[Image Missing: {img_path}]")

def generate_document():
    doc = Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.header_distance = Inches(0.3)
        section.footer_distance = Inches(0.3)
    
    base_path = r"c:\Users\HP\Desktop\Expensive categerization"
    img_dir = os.path.join(base_path, "Images")
    bend_dir = os.path.join(img_dir, "BEND")
    fend_dir = os.path.join(img_dir, "FEND")
    res_dir = os.path.join(img_dir, "Results")
    
    left_logo = os.path.join(img_dir, "leftsidelogo.jpeg")
    right_logo = os.path.join(img_dir, "rightsidelogo.jpeg")
    
    add_header(doc, left_logo, right_logo)
    
    # --- PAGE 1: Title & Description ---
    add_page_title(doc, "Expense Categorization Tool", "AI-Powered Financial Insights and Analysis")
    doc.add_paragraph()
    add_heading(doc, "Project Description", level=1)
    doc.add_paragraph(
        "The Expense Categorization Tool is an intelligent financial management platform that leverages advanced AI "
        "to provide comprehensive expense analysis and categorization. The platform addresses the challenge of "
        "unorganized financial records by delivering AI-powered insights, detailed spending breakdowns, and "
        "visual reports to improve financial awareness and budgeting habits.\n\n"
        "Using cutting-edge AI models, the system analyzes transaction data from CSV files or Google Sheets, "
        "automatically classifying expenses into meaningful categories like Food, Travel, Shopping, and Utilities. "
        "The system ensures high accuracy through sophisticated AI analysis and data processing algorithms, "
        "allowing users to identify spending patterns and optimize their financial health.\n\n"
        "The Expense Categorization Tool transforms financial tracking into an intelligent, user-friendly experience "
        "through its modern interface, comprehensive feature set, and AI-powered analysis that provides personalized "
        "budgeting guidance. By automating the tedious task of manual categorization, the tool helps users make "
        "smarter financial decisions and maintain better control over their expenses."
    )
    doc.add_page_break()
    
    # --- PAGE 2: Scenarios ---
    add_heading(doc, "Scenarios", level=1)
    scenarios = [
        ("Scenario 1 – Personal Budgeting:", 
         "A university graduate starting their first job wants to manage their salary effectively but finds manual tracking overwhelming. "
         "By uploading their monthly bank statements to the Expense Categorization Tool, the AI automatically groups "
         "transactions into categories like 'Rent', 'Groceries', 'Dining Out', and 'Transport'. The user receives a visual "
         "breakdown that highlights excessive spending in 'Entertainment', enabling them to adjust their habits and "
         "meet their savings goals for the month."),
        ("Scenario 2 – Small Business Expense Tracking:", 
         "A small retail business owner needs to organize hundreds of supplier payments and utility bills for annual tax filing. "
         "Instead of spending days on manual entry, they upload their transaction history. The tool identifies and "
         "categorizes 'Inventory Purchases', 'Utilities', 'Marketing', and 'Payroll'. This provides a clear financial "
         "overview for their accountant, ensuring maximum tax deductions and compliance with financial regulations."),
        ("Scenario 3 – Freelance Travel Management:", 
         "A freelance consultant travels frequently for various clients and needs to separate billable expenses. "
         "The tool categorizes trip-related costs like 'Airfare', 'Hotel stays', and 'Client Meals'. The consultant "
         "can then generate specific reports for each client, streamlining the reimbursement process and "
         "maintaining professional financial records without the hassle of manual spreadsheet management."),
        ("Scenario 4 – Family Financial Oversight:", 
         "A family of four uses the tool to monitor their shared household expenses. The AI identifies recurring "
         "subscription costs and highlights a significant increase in 'Utilities' compared to the previous month. "
         "This insight prompts the family to investigate a potential water leak and cancel unused streaming services, "
         "leading to substantial annual savings and better overall financial planning.")
    ]
    for title, desc in scenarios:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        run.font.size = Pt(12)
        doc.add_paragraph(desc)
    doc.add_page_break()
    
    # --- PAGE 3: Architecture Overview ---
    add_heading(doc, "Architecture Overview", level=1)
    doc.add_paragraph(
        "The Expense Categorization Tool is built as a modular platform combining a robust Python/Flask backend with "
        "a modern React/Vite frontend. The architecture prioritizes data integrity, processing efficiency, and "
        "a seamless user experience by leveraging AI-driven categorization and dynamic data visualization."
    )
    
    add_heading(doc, "Core Technologies", level=2)
    techs = [
        "Python & Flask: Lightweight and efficient backend for routing and API management.",
        "AI Model Integration: Advanced natural language processing for intelligent transaction classification.",
        "React.js: High-performance frontend library for building interactive user interfaces.",
        "Vite: Ultra-fast build tool and development server for modern web projects.",
        "Pandas & NumPy: Powerful data analysis libraries for processing complex CSV structures.",
        "CSS & Tailwind: Modern styling for a responsive and premium visual experience.",
        "Chart.js/Matplotlib: Dynamic charting for visual financial reports and insights."
    ]
    for tech in techs:
        doc.add_paragraph(tech, style='List Bullet')
        
    add_heading(doc, "Pre-requisites", level=1)
    add_heading(doc, "Software Requirements", level=2)
    softs = [
        "Python 3.8+: Core programming language for backend development.",
        "Node.js & npm: Essential for managing frontend dependencies and build tools.",
        "Git: Version control for tracking project development and collaboration.",
        "OpenAI API Key: Required for accessing advanced AI categorization models.",
        "Postman: For testing and validating backend API endpoints.",
        "Modern Web Browser: For interacting with the frontend application."
    ]
    for s in softs:
        doc.add_paragraph(s, style='List Bullet')
        
    add_heading(doc, "Knowledge Prerequisites", level=2)
    knows = [
        "Python Programming: Proficiency in Flask and data handling libraries.",
        "JavaScript & React: Understanding of component-based UI development.",
        "API Integration: Experience with RESTful services and AI model interaction.",
        "Financial Literacy: Basic understanding of expense categories and budgeting.",
        "Data Management: Skills in handling and cleaning CSV/Excel data."
    ]
    for k in knows:
        doc.add_paragraph(k, style='List Bullet')
        
    add_heading(doc, "Hardware Requirements", level=2)
    hards = [
        "Processor: Dual-core 2.4GHz or better (Intel i5/AMD Ryzen 5 recommended).",
        "RAM: 8GB minimum (16GB recommended for smooth development).",
        "Storage: 10GB free space for project files and dependencies.",
        "Internet Connection: Stable connection for API calls and package management."
    ]
    for h in hards:
        doc.add_paragraph(h, style='List Bullet')
    doc.add_page_break()
    
    # --- PAGE 4-6: Project Workflow ---
    add_heading(doc, "Project Workflow", level=1)
    
    add_heading(doc, "Phase 1: Environment Setup & Project Initialization", level=2)
    doc.add_paragraph("Establish the development infrastructure and configure core services.")
    p1_acts = [
        ("Activity 1.1: Set up Development Environment", [
            "Initialize Python virtual environment and install backend dependencies.",
            "Set up Node.js environment and initialize Vite project for the frontend.",
            "Configure environment variables and secure API key storage."
        ]),
        ("Activity 1.2: Project Architecture Design", [
            "Define the backend directory structure and Flask application entry point.",
            "Design the frontend component hierarchy and state management strategy.",
            "Establish Git repository and branch management workflow."
        ])
    ]
    for title, bullets in p1_acts:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet 2')

    add_heading(doc, "Phase 2: Backend & AI Engine Development", level=2)
    doc.add_paragraph("Build the core processing logic and integrate AI models.")
    p2_acts = [
        ("Activity 2.1: Implement Data Upload & Parsing", [
            "Create Flask routes for secure CSV and Excel file uploads.",
            "Develop data cleaning scripts using Pandas to handle inconsistent formats.",
            "Implement error handling for malformed or missing transaction data."
        ]),
        ("Activity 2.2: AI Categorization Logic", [
            "Design prompts for the AI model to accurately classify transaction descriptions.",
            "Implement API calling logic with retry mechanisms and timeout handling.",
            "Build a fallback rule-based system for common or high-confidence categories."
        ])
    ]
    for title, bullets in p2_acts:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet 2')
    doc.add_page_break()

    add_heading(doc, "Phase 3: Frontend & Dashboard Development", level=2)
    doc.add_paragraph("Create a responsive user interface and interactive visualizations.")
    p3_acts = [
        ("Activity 3.1: Design Core UI Components", [
            "Develop the transaction list view with filtering and sorting capabilities.",
            "Build the file upload interface with real-time progress indicators.",
            "Create user authentication screens for secure data access."
        ]),
        ("Activity 3.2: Interactive Data Visualization", [
            "Integrate Chart.js to display category-wise spending breakdowns.",
            "Develop monthly trend charts and budget comparison widgets.",
            "Implement interactive tooltips and detailed drill-down reports."
        ])
    ]
    for title, bullets in p3_acts:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet 2')

    add_heading(doc, "Phase 4: Integration & System Testing", level=2)
    doc.add_paragraph("Ensure seamless communication between components and validate accuracy.")
    p4_acts = [
        ("Activity 4.1: End-to-End Testing", [
            "Verify the full flow from file upload to AI categorization and UI display.",
            "Conduct cross-browser compatibility testing for consistent experience.",
            "Perform load testing with large datasets to ensure system stability."
        ])
    ]
    for title, bullets in p4_acts:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet 2')

    add_heading(doc, "Phase 5: Deployment & Final Documentation", level=2)
    p5_acts = [
        ("Activity 5.1: Final Deployment & Review", [
            "Deploy the application to a cloud provider or local production server.",
            "Perform final security audits and API performance tuning.",
            "Generate comprehensive user manuals and technical project reports."
        ])
    ]
    for title, bullets in p5_acts:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.font.bold = True
        for b in bullets:
            doc.add_paragraph(b, style='List Bullet 2')
    doc.add_page_break()
    
    # --- PAGE 7: Technical Architecture ---
    add_heading(doc, "Technical Architecture", level=1)
    doc.add_paragraph("The technical architecture of the Expense Categorization Tool follows a modern decoupled approach.")
    add_heading(doc, "Project Structure", level=2)
    add_image_with_caption(doc, os.path.join(fend_dir, "p-structure.jpeg"), "Detailed Project Directory Structure")
    doc.add_paragraph(
        "The project structure is meticulously organized into backend and frontend directories. "
        "The 'backend' folder contains the Flask application, API routes, and data processing logic. "
        "The 'frontend' folder houses the React components, assets, and styling configurations. "
        "This separation ensures a scalable and maintainable codebase, allowing for independent development "
        "and deployment of both components."
    )
    doc.add_page_break()
    
    # --- MILESTONES ---
    
    # Milestone 1
    add_heading(doc, "MILESTONE 1: Environment Setup & Requirements", level=1)
    doc.add_paragraph("This initial milestone focuses on defining the project scope and setting up the development environment.")
    add_image_with_caption(doc, os.path.join(img_dir, "req.jpeg"), "Activity 1.1: Project Requirements Gathering")
    doc.add_paragraph(
        "The requirements document outlines the functional and non-functional specifications of the tool. "
        "It includes details on data privacy, AI model selection, and the desired user experience milestones."
    )
    doc.add_page_break()
    add_image_with_caption(doc, os.path.join(img_dir, "package install.jpeg"), "Activity 1.2: Dependency Installation")
    doc.add_paragraph(
        "Installing the core libraries such as Flask for the backend and React for the frontend. "
        "All dependencies are tracked in requirements.txt and package.json to ensure environment consistency."
    )
    doc.add_page_break()
    
    # Milestone 2: BEND
    add_heading(doc, "MILESTONE 2: Core Backend Development", level=1)
    doc.add_paragraph("Development of the backend infrastructure and AI processing engine.")
    
    bend_descriptions = {
        "app2.jpeg": "Backend API Configuration: Setting up the Flask application and defining the initial API routes for data handling.",
        "app3.jpeg": "AI Categorization Engine: Implementation of the core logic that communicates with the AI model to classify expenses.",
        "app4.jpeg": "Data Processing Logic: Developing the Pandas-based engine for cleaning and structuring uploaded transaction data."
    }
    
    for img, desc in bend_descriptions.items():
        add_image_with_caption(doc, os.path.join(bend_dir, img), f"Backend Activity: {img}")
        doc.add_paragraph(desc)
        doc.add_paragraph(
            "This component is critical for ensuring that the application can handle various data formats "
            "and provide accurate, real-time results to the end user. The backend serves as the bridge between "
            "the raw data and the intelligent insights provided by the AI model."
        )
        doc.add_page_break()
        
    # Milestone 3: FEND
    add_heading(doc, "MILESTONE 3: Frontend Development", level=1)
    doc.add_paragraph("Designing and building the interactive user interface for the Expense Categorization tool.")
    
    fend_imgs = ["index.jpeg", "main1.jpeg", "main2.jpeg", "main3.jpeg", "main4.jpeg", "react.jpeg", "sample Dataset.jpeg"]
    fend_descriptions = {
        "index.jpeg": "Main Application Entry: The entry point of the React application, managing global state and routing.",
        "main1.jpeg": "Dashboard Component: Implementation of the primary dashboard view with summary widgets.",
        "main2.jpeg": "Transaction List View: Building the interactive table for displaying and filtering transactions.",
        "main3.jpeg": "Upload Interface: Creating the user-friendly drag-and-drop interface for financial data files.",
        "main4.jpeg": "AI Insights Display: Developing the component that renders intelligent financial advice from the AI model.",
        "react.jpeg": "React Component Architecture: Overview of the modular component structure used in the frontend.",
        "sample Dataset.jpeg": "Data Format Specification: Defining the expected CSV structure for consistent processing."
    }
    
    for img in fend_imgs:
        add_image_with_caption(doc, os.path.join(fend_dir, img), f"Frontend Activity: {img}")
        doc.add_paragraph(fend_descriptions.get(img, "Frontend implementation detail."))
        doc.add_paragraph(
            "The frontend is designed with a focus on usability and clarity. By using a modular approach, "
            "each component remains independent and easy to update. The interface ensures that users can "
            "easily navigate through their financial data and understand the AI-generated insights."
        )
        doc.add_page_break()
        
    # Milestone 4: Deployment & Results
    add_heading(doc, "MILESTONE 4: Deployment & Final Results", level=1)
    doc.add_paragraph("Deploying the application and verifying the final output of the categorization tool.")
    
    add_heading(doc, "Deployment Environment Setup", level=2)
    term1 = os.path.join(fend_dir, "terminal1.jpeg")
    term2 = os.path.join(fend_dir, "terminal2.jpeg")
    
    add_image_with_caption(doc, term1, "Activity 4.1: Backend Server Deployment")
    doc.add_paragraph("Verification of the Flask server starting up and listening for incoming API requests.")
    doc.add_page_break()
    
    add_image_with_caption(doc, term2, "Activity 4.2: Frontend Development Server")
    doc.add_paragraph("Confirming the Vite development server is running and the frontend is accessible via the browser.")
    doc.add_page_break()
    
    add_heading(doc, "Final Application Results", level=2)
    res_imgs = [
        "Home page.jpeg", "login.jpeg", "signup.jpeg", "dashboard.jpeg", 
        "Upload.jpeg", "after file upload.jpeg", "transactions.jpeg", 
        "AI insights.jpeg", "profile.jpeg", "app.jpeg"
    ]
    res_descriptions = {
        "Home page.jpeg": "Application Landing Page: The professional entry point for users to access the tool.",
        "login.jpeg": "Secure Login Interface: Authentication screen ensuring user data privacy and security.",
        "signup.jpeg": "User Registration: New user onboarding process for personalized financial tracking.",
        "dashboard.jpeg": "Main Dashboard Overview: Visual summary of categorized expenses and spending trends.",
        "Upload.jpeg": "File Upload Portal: Easy interface for importing transaction data from external files.",
        "after file upload.jpeg": "Data Processing Confirmation: Real-time feedback after successfully parsing transaction files.",
        "transactions.jpeg": "Detailed Transaction View: Comprehensive list of categorized transactions with manual override options.",
        "AI insights.jpeg": "Intelligent Financial Reports: AI-generated advice and detailed spending analysis.",
        "profile.jpeg": "User Profile Management: Settings for managing account details and financial preferences.",
        "app.jpeg": "Full Application View: Integrated view showing the seamless connection between all modules."
    }
    
    for img in res_imgs:
        add_image_with_caption(doc, os.path.join(res_dir, img), f"Result: {img}")
        doc.add_paragraph(res_descriptions.get(img, "Application result screenshot."))
        doc.add_paragraph(
            "The final results demonstrate the tool's effectiveness in transforming raw transaction data "
            "into clear, categorized financial insights. The combination of accurate AI classification "
            "and intuitive visualization provides users with a powerful tool for personal finance management."
        )
        doc.add_page_break()
        
    # --- Conclusion ---
    add_heading(doc, "Conclusion", level=1)
    doc.add_paragraph(
        "The Expense Categorization Tool successfully provides an automated, AI-driven solution for managing "
        "complex financial records. By integrating a robust Flask backend with a modern React frontend, "
        "the application delivers high accuracy in transaction classification and a premium user experience.\n\n"
        "The project milestones demonstrate the systematic development process, from environment setup "
        "to final deployment and result validation. This tool serves as a powerful example of how AI can "
        "be applied to everyday financial tasks, helping users gain better control over their spending "
        "and build a more secure financial future."
    )
    
    # Save the document
    output_file = os.path.join(base_path, "Expense_Categorization_Final_Doc.docx")
    doc.save(output_file)
    print(f"Document generated successfully: {output_file}")


if __name__ == "__main__":
    generate_document()
